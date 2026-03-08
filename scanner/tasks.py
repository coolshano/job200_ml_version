from celery import shared_task
from django.core.mail import EmailMultiAlternatives
from django.template.loader import render_to_string
from django.conf import settings
from .models import Resume
from .utils import (
    extract_resume_text,
    calculate_ats_score,
    find_missing_keywords,
    generate_recommendations
)
import json
#from .parser import ResumeParser
from django.core.mail import EmailMessage
from .cv_generator import generate_and_optimize_cv

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate
from reportlab.platypus import Paragraph
from reportlab.platypus import Spacer
from reportlab.platypus import ListFlowable, ListItem
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import A4
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from django.template.loader import render_to_string
from .cv_generator import generate_and_optimize_cv

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .word_generator import create_word_resume

@shared_task
def scan_resume_task(resume_id):

    resume = Resume.objects.get(id=resume_id)

    try:

        # ============================================
        # Stage 1 — Extract Text
        # ============================================

        resume.status = "extracting_text"
        resume.save(update_fields=["status"])

        resume.resume_file.open()
        text = extract_resume_text(resume.resume_file)

        # ============================================
        # Stage 2 — Structured Parsing
        # ============================================

        resume.status = "parsing_resume"
        resume.save(update_fields=["status"])

        #parser = ResumeParser(text)
        #structured_data = parser.parse()

        #resume.parsed_data = structured_data
        #resume.save(update_fields=["parsed_data"])

        # ============================================
        # Stage 3 — Keyword Matching
        # ============================================

        resume.status = "matching_keywords"
        resume.save(update_fields=["status"])

        missing_keywords = find_missing_keywords(
            text,
            resume.job_description
        )

        # ============================================
        # Stage 4 — Advanced Scoring
        # ============================================

        resume.status = "calculating_score"
        resume.save(update_fields=["status"])

        score_data = calculate_ats_score(
            text,
            resume.job_description
        )

        if isinstance(score_data, dict):
            score = score_data.get("total_score", 0)
            breakdown = score_data
        else:
            score = score_data
            breakdown = {
                "total_score": score
            }

        # ============================================
        # Generate Recommendations
        # ============================================

        recommendations = generate_recommendations(
            text,
            resume.job_description,
            score
        )

        # 🔍 DEBUG START
        print("===== DEBUG RECOMMENDATIONS =====")
        print("Score:", score)
        print("Generated recommendations:", recommendations)
        print("Breakdown BEFORE adding recs:", breakdown)
        # 🔍 DEBUG END

        breakdown["recommendations"] = recommendations

        # ============================================
        # Save Results
        # ============================================

        resume.score = float(round(score, 2))
        resume.missing_keywords = missing_keywords
        resume.score_breakdown = json.loads(json.dumps(breakdown, default=float))
        resume.status = "generating_optimized_resume"

        resume.save(update_fields=[
            "score",
            "missing_keywords",
            "score_breakdown",
            "status"
        ])

        # 🔍 DEBUG SAVE CHECK
        resume.refresh_from_db()
        print("Breakdown AFTER save:", resume.score_breakdown)
        print("=================================")

        # ============================================
        # Stage 5 — Generate Optimized Resume
        # ============================================

        #generate_ats_docx(resume)

        resume.status = "completed"
        resume.save(update_fields=["status"])

        # ============================================
        # Candidate Email
        # ============================================

        html_content = render_to_string(
            "emails/ats_result.html",
            {
                "name": resume.name,
                "score": resume.score,
                "missing_keywords": missing_keywords,
                "score_breakdown": breakdown,
                "recommendations": recommendations,
            },
        )

        email = EmailMultiAlternatives(
            subject="Your ATS Resume Scan Result",
            body="Your resume scan results are ready.",
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=[resume.email],
        )

        email.attach_alternative(html_content, "text/html")
        email.send()

        # ============================================
        # Recruiter Alert
        # ============================================

        if resume.score >= settings.ATS_ALERT_THRESHOLD:

            recruiter_email = EmailMultiAlternatives(
                subject="High-Match Candidate Alert",
                body=f"""
High-match candidate detected

Name: {resume.name}
Email: {resume.email}
Job Title: {resume.job_title}
ATS Score: {resume.score}%
                """,
                from_email=settings.DEFAULT_FROM_EMAIL,
                to=settings.RECRUITER_NOTIFICATION_EMAILS,
            )

            resume.resume_file.open()
            file_content = resume.resume_file.read()

            recruiter_email.attach(
                resume.resume_file.name.split("/")[-1],
                file_content,
                "application/pdf",
            )

            recruiter_email.send()

    except Exception as e:
        resume.status = "failed"
        resume.save(update_fields=["status"])
        raise e
    

@shared_task
def generate_and_email_cv(data, job_description, user_email):

    # 1️⃣ Generate Resume (JSON structure)
    resume_data, score = generate_and_optimize_cv(data, job_description)

    # 2️⃣ Create Word Document
    document = Document()

    title = document.add_heading(data.get("name", "Resume"), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_paragraph(f"ATS Optimization Score: {score}%")

    # Summary
    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(resume_data.get("summary", ""))

    # Experience
    document.add_heading("Experience", level=1)
    for exp in resume_data.get("experience", []):
        document.add_paragraph(exp, style="List Bullet")

    # Skills
    document.add_heading("Skills", level=1)
    for skill in resume_data.get("skills", []):
        document.add_paragraph(skill, style="List Bullet")

    # Education
    document.add_heading("Education", level=1)
    document.add_paragraph(resume_data.get("education", ""))

    # Certifications
    document.add_heading("Certifications", level=1)
    for cert in resume_data.get("certifications", []):
        document.add_paragraph(cert, style="List Bullet")

    # Projects
    document.add_heading("Projects", level=1)
    for project in resume_data.get("projects", []):
        document.add_paragraph(project, style="List Bullet")

    # Achievements
    document.add_heading("Achievements", level=1)
    for ach in resume_data.get("achievements", []):
        document.add_paragraph(ach, style="List Bullet")

    # 3️⃣ Save Word file
    safe_name = data.get("name", "resume").replace(" ", "_")
    file_name = f"{safe_name}_resume.docx"
    file_path = os.path.join("/tmp", file_name)

    document.save(file_path)

    # 4️⃣ Render Email Template
    html_content = render_to_string(
        "emails/resume_generated.html",
        {
            "name": data.get("name"),
            "score": score,
        }
    )

    # 5️⃣ Send Email
    email = EmailMessage(
        subject="Your AI Resume is Ready - JOB200 ATS",
        body=html_content,
        to=[user_email],
    )

    email.content_subtype = "html"
    email.attach_file(file_path)

    email.send()

    # 6️⃣ Clean up temp file
    if os.path.exists(file_path):
        os.remove(file_path)

    return "Resume Word Document Sent Successfully"