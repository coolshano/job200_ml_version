from docx import Document
from docx.shared import Pt


def create_word_resume(resume_data, name, score):

    document = Document()

    document.add_heading(name, level=0)

    document.add_paragraph(f"ATS Optimization Score: {score}%")

    document.add_heading("Professional Summary", level=1)
    document.add_paragraph(resume_data.get("summary", ""))

    document.add_heading("Experience", level=1)
    for exp in resume_data.get("experience", []):
        document.add_paragraph(exp, style="List Bullet")

    document.add_heading("Skills", level=1)
    for skill in resume_data.get("skills", []):
        document.add_paragraph(skill, style="List Bullet")

    document.add_heading("Education", level=1)
    document.add_paragraph(resume_data.get("education", ""))

    document.add_heading("Certifications", level=1)
    for cert in resume_data.get("certifications", []):
        document.add_paragraph(cert, style="List Bullet")

    document.add_heading("Projects", level=1)
    for project in resume_data.get("projects", []):
        document.add_paragraph(project, style="List Bullet")

    document.add_heading("Achievements", level=1)
    for ach in resume_data.get("achievements", []):
        document.add_paragraph(ach, style="List Bullet")

    file_path = f"/tmp/{name.replace(' ','_')}_resume.docx"

    document.save(file_path)

    return file_path